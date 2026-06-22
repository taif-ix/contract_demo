from pathlib import Path
import textract


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_text(file_path: Path) -> str:
    """
    Extract plain text from a file based on its extension.
    Supports: PDF, DOCX, DOC, TXT
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_from_pdf(file_path)
    elif suffix == ".docx":
        return _extract_from_docx(file_path)
    elif suffix == ".doc":
        return _extract_from_doc(file_path)
    elif suffix == ".txt":
        return _extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_from_pdf(file_path: Path) -> str:
    import fitz
    text_output = ""
    with fitz.open(file_path) as doc:
        for page_no, page in enumerate(doc, start=1):
            text_output += f"\n--- Page {page_no} ---\n{page.get_text()}"
    return text_output.strip()


def _extract_from_docx(file_path: Path) -> str:
    from docx import Document
    doc = Document(str(file_path))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs)


def _extract_from_doc(file_path: Path) -> str:
    """
    Extract old Microsoft Word .doc files.
    """

    import subprocess

    # Try antiword
    try:
        result = subprocess.run(
            ["antiword", str(file_path)],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode == 0:
            text = result.stdout.strip()

            if text:
                return text

    except Exception:
        pass


    # Try textract
    try:
        import textract

        text = textract.process(
            str(file_path)
        ).decode(
            "utf-8",
            errors="ignore"
        )

        if text.strip():
            return text.strip()

    except Exception:
        pass


    return ""


def _extract_from_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore").strip()
